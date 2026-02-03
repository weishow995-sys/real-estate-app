import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime
import io
import docx

# --- 1. 網頁介面大字體設定 ---
st.set_page_config(page_title="房地產終極評估系統", layout="centered")
st.title("🏠 房地產一鍵評估系統 (V5)")
st.write("請上傳謄本照片或 PDF，系統將自動生成 Word 報告。")

# --- 2. 您的 API KEY (請確認引號內有貼上那一串 AIza... 的代碼) ---
API_KEY = "您的_API_KEY_貼在這邊"

# --- 3. Word 排版輔助工具 ---
def set_font(run, size=14, bold=False, color=None):
    run.font.name = 'Microsoft JhengHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000FF')
    rPr.append(c)
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)
    f = docx.oxml.shared.OxmlElement('w:rFonts')
    f.set(docx.oxml.shared.qn('w:eastAsia'), 'Microsoft JhengHei')
    rPr.append(f)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# --- 4. 檔案上傳介面 ---
uploaded_file = st.file_uploader("選擇檔案", type=["pdf", "png", "jpg", "jpeg"])

if uploaded_file and API_KEY != "您的_API_KEY_貼在這邊":
    if st.button("🚀 點此開始產出報告"):
        with st.spinner("AI 正在深度解析並計算殘值..."):
            try:
                genai.configure(api_key=API_KEY)
                # 使用最穩定的 flash 模型
                model = genai.GenerativeModel('gemini-1.5-flash')
                
                # 強制指令：包含身分證英文字母、街景、殘值試算
                prompt = """
                請解析此謄本，並產出以下資訊：
                1. 產權警示：查封/限制登記/民間二胎。
                2. 社區建築：社區名、構造、樓層、屋齡。
                3. 所有權人：姓名、完整身分證(必須包含首位英文字母，如 R220*****9)、持分、戶籍地、地址。
                4. 貸款殘值：列出銀行、設定額、日期。採30年2.15%利率試算目前餘額。
                5. 二胎空間：計算(市場80%價格 - 剩餘貸款)。
                6. 交通：到國道與火車站車程。
                禁止出現任何 標記。
                """
                
                mime_type = "application/pdf" if uploaded_file.name.lower().endswith(".pdf") else uploaded_file.type
                response = model.generate_content([prompt, {"mime_type": mime_type, "data": uploaded_file.getvalue()}])
                
                # --- 製作 Word 檔案 ---
                doc = Document()
                title = doc.add_heading('', 0)
                run_t = title.add_run('房地產全方位終極評估報告書')
                set_font(run_t, size=20, bold=True, color=RGBColor(0, 51, 153))
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 內容文字 (AI 回傳的結果)
                p = doc.add_paragraph()
                set_font(p.add_run(response.text), size=14)
                
                # 強制加入可點擊超連結 (以地址為基礎)
                doc.add_heading('', level=1).add_run('外部資源連結').font.size = Pt(16)
                p_link = doc.add_paragraph()
                set_font(p_link.add_run("Google 街景圖連結："))
                # 這裡預留一個連結位置
                add_hyperlink(p_link, "點此開啟 Google 街景", "https://www.google.com/maps")
                
                # 產出檔案
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                
                st.success("報告生成成功！")
                st.download_button(
                    label="📥 下載 Word 報告書",
                    data=buf,
                    file_name=f"房產評估報告_{datetime.date.today()}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"發生錯誤：{e}")
elif not uploaded_file:
    st.info("請上傳檔案後點擊開始。")
